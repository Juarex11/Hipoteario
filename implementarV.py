import datetime
from num2words import num2words
from docx import Document
from docx.shared import Pt
import json
import os

meses_en_espanol = {
    1: "ENERO",
    2: "FEBRERO",
    3: "MARZO",
    4: "ABRIL",
    5: "MAYO",
    6: "JUNIO",
    7: "JULIO",
    8: "AGOSTO",
    9: "SEPTIEMBRE",
    10: "OCTUBRE",
    11: "NOVIEMBRE",
    12: "DICIEMBRE"
}

def leer_json(ruta):
    with open(ruta, 'r', encoding='utf-8') as archivo_json:
        datos = json.load(archivo_json)
    return datos

def eliminar_espacios_finales(texto):
    # Eliminar espacios en blanco al final del texto
    return ''.join(texto.split())

def agregar_representado_por(registro, frase_generada):
    representantes = registro.get("Representantes", [])

    # Verificar si hay representantes
    if representantes:
        frase_generada += " REPRESENTADO POR"

    return frase_generada

def generar_detalles(registro):
    detalles = []

    # Verificar cada campo y agregarlo solo si no está vacío o nulo
    #if registro["Id"]:
    #    detalles.append(f'id: {registro["Id"]}')
    #if registro["Tipo"]:
    #    detalles.append(f'Tipo: {registro["Tipo"]}')
    if registro["Nombre"]:
        detalles.append(f'{registro["Nombre"]}')
    if registro["Nacionalidad"]:
        detalles.append(f'Nacionalidad: {registro["Nacionalidad"]}')
    if registro["Numero_de_documento"]:
        detalles.append(f' {registro["Tipo_de_documento"]} {registro["Numero_de_documento"]}')
    if registro["Estado_civil"]:
        detalles.append(f'Estado civil: {registro["Estado_civil"]}')
    if registro["Domicilio"]:
        detalles.append(f'Domicilio: {registro["Domicilio"]}')
    
    # Unir los detalles en una cadena
    resultado = ', '.join(detalles)

    # Llamar a la función para agregar "REPRESENTADO POR" si es necesario
    resultado = agregar_representado_por(registro, resultado)

    # Devolver None si no hay detalles para evitar agregar un campo vacío en el documento Word
    return resultado if resultado else None

def agregar_espaciado_entre_palabras(texto):
    # Agregar un espacio entre cada palabra
    return ' '.join(texto.split())

def cambiar_formato_texto(parrafo, fuente='Arial Narrow', tamaño=9):
    run = parrafo.runs[0]  # Suponiendo que el párrafo ya tiene al menos un run
    font = run.font
    font.name = fuente
    font.size = Pt(tamaño)
    # Convertir el texto a mayúsculas
    run.text = run.text.upper()

def main():
    # Leer datos del archivo JSON
    datos_json = leer_json('json/archivo.json')

    # Leer la plantilla del documento Word
    plantilla_doc = Document('plantilla/MODELO PLANTILLA.docx')

    # Obtener la fecha actual
    fecha_actual = datetime.datetime.now()
    # Obtener el día en letras
    dia_en_letras = num2words(fecha_actual.day, lang='es').upper()
    # Obtener el nombre del mes en español
    mes_en_espanol = meses_en_espanol[fecha_actual.month]
    # Obtener el año en letras
    anio_en_letras = num2words(fecha_actual.year, lang='es').upper()
    # Crear el texto de la fecha
    texto_fecha = f"EL {dia_en_letras} DE {mes_en_espanol} DEL AÑO {anio_en_letras}"

    # Inicializar una lista para almacenar todos los detalles
    detalles = []

    # Iterar sobre todos los registros en el archivo JSON y agregar los detalles a la lista
    for registro in datos_json:
        detalle = generar_detalles(registro)
        if detalle is not None:
            detalles.append(detalle)

    # Pegar todos los detalles en un solo párrafo
    resultado = '\n'.join(detalles)

    # Eliminar espacios en blanco al final del texto generado
    resultado = agregar_espaciado_entre_palabras(resultado)

    # Definir la cadena a reemplazar
    cadena_a_reemplazar = 'NOMBRE, PERUANO/A, SOLTERO/CASADO, OCUPACION …, CON DOCUMENTO NACIONAL DE IDENTIDAD …., CON DOMICILIO EN …..(NOTA VERIFICAR QUE SEA EL DEL D.N.I.; SINO INDICAR QUE EL DECLARA QUE SU DOMICILIO ACTUAL ES …. ); Y DE LA OTRA PARTE: NOMBRE, PERUANO/A, SOLTERO/CASADO, OCUPACION …, CON DOCUMENTO NACIONAL DE IDENTIDAD …., CON DOMICILIO EN …..'

    # Reemplazar las palabras en la plantilla con la lista de detalles
    for p in plantilla_doc.paragraphs:
        if cadena_a_reemplazar in p.text:
            p.text = p.text.replace(cadena_a_reemplazar, resultado)
            cambiar_formato_texto(p)

    # Reemplazar el texto en la plantilla con la fecha actual
    cadena_a_reemplazar_fecha = 'EL XXXXXXXXXXXXXX DE XXXXXXXXXXXXXX DEL AÑO DOS MIL XXXXXXXXXXXXXX'
    for p in plantilla_doc.paragraphs:
        if cadena_a_reemplazar_fecha in p.text:
            p.text = p.text.replace(cadena_a_reemplazar_fecha, texto_fecha)
            cambiar_formato_texto(p)

    # Guardar el documento modificado en la carpeta 'modificado' con el nombre 'ejemplo_terminado.docx'
    if not os.path.exists('modificado'):
        os.makedirs('modificado')
    plantilla_doc.save('modificado/ejemplo_terminado.docx')
    print("Programa finalizado")

if __name__ == "__main__":
    main()
