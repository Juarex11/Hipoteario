import docx
import openai
import os
import time
<<<<<<< HEAD
from docx.shared import Pt

# Cambiamos formato de texyo usando una funcion de otro archivo
from implementar import cambiar_formato_texto

# Configura la API de OpenAI con tu clave
openai.api_key = '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'
=======
from implementar import cambiar_formato_texto
from docx.shared import Pt


# Configura la API de OpenAI con tu clave
openai.api_key = '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'
>>>>>>> 07757c7e23417c855c84eda7f4a847e73c9a8d83

# Función para obtener los primeros 3 párrafos de un archivo DOCX
def obtener_primeros_parrafos(docx_path, num_parrafos=10):
    doc = docx.Document(docx_path)
    parrafos = [p.text for p in doc.paragraphs[:num_parrafos]]
    return "\n".join(parrafos)

# Función para generar texto con la IA
def generar_texto_ia(prompt):
    # Llamar a la API de GPT-3 para completar el prompt
<<<<<<< HEAD
    response = openai.chat.completions.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
                    ]
    )
    return response.choices[0].message.content.strip()
=======
    response = openai.Completion.create(
        model="text-davinci-003",
        prompt=prompt[:4096],  # Reduzco la longitud del prompt a 4096 caracteres
        max_tokens=100,
        temperature=0.7
    )

    return response.choices[0].text.strip()
>>>>>>> 07757c7e23417c855c84eda7f4a847e73c9a8d83

def cambiar_formato_texto(parrafo, fuente='Arial Narrow', tamaño=9):
    run = parrafo.runs[0]
    font = run.font
    font.name = fuente
    font.size = Pt(tamaño)
    run.text = run.text.upper()

def main():
    # Ruta del archivo DOCX original
    carpeta_documentos = 'documentos'
<<<<<<< HEAD
    archivo_docx_original = 'IBR.docx'
=======
    archivo_docx_original = 'BCP - PYME.docx'
>>>>>>> 07757c7e23417c855c84eda7f4a847e73c9a8d83
    ruta_completa_docx_original = os.path.join(carpeta_documentos, archivo_docx_original)

    # Verificar la existencia del archivo original
    if not os.path.exists(ruta_completa_docx_original):
        print(f"El archivo {ruta_completa_docx_original} no existe. Verifica la ruta y el nombre del archivo.")
        exit()

    # Obtener los primeros 3 párrafos del archivo DOCX original
    primeros_parrafos = obtener_primeros_parrafos(ruta_completa_docx_original)

    # Prompt más específico para la IA
    prompt_ia = f"Dame este texto según lo entendido (identifica la acción, ejemplo compra venta de inmueble), que otorga (nombre del otorgante negocio) a favor de (nombre de la persona a favor) que otorga a (nombre del banco) a favor de (nuevamente el nombre de la persona a favor, IMPORTANTE REPETIR) redactar bien: {primeros_parrafos} ESCRITURA PUBLICA DE "

    # Generar texto con la IA
    texto_generado = generar_texto_ia(prompt_ia)

    # Mostrar mensaje mientras se espera la respuesta de la IA
    print("Esperando respuesta de la IA...")

    # Esperar unos segundos antes de continuar (opcional)
    time.sleep(5)

    # Ruta del archivo DOCX modificado
    carpeta_modificado = 'modificado'
    archivo_docx_modificado = 'ejemplo_terminado.docx'
    ruta_completa_docx_modificado = os.path.join(carpeta_modificado, archivo_docx_modificado)

    # Verificar la existencia del archivo modificado
    if not os.path.exists(ruta_completa_docx_modificado):
        print(f"El archivo {ruta_completa_docx_modificado} no existe. Verifica la ruta y el nombre del archivo.")
        exit()

    # Cargar el documento Word modificado
    doc_modificado = docx.Document(ruta_completa_docx_modificado)

    # Reemplazar el texto en el documento modificado
    texto_viejo = "ACTO ; QUE OTORGAN NOMBRE; A FAVOR DE NOMBRE"
    for p in doc_modificado.paragraphs:
        if texto_viejo in p.text:
         p.text = p.text.replace(texto_viejo, texto_generado)
        # Cambiar el formato del párrafo, no del texto generado
         cambiar_formato_texto(p)

    # Guardar el documento Word modificado
    doc_modificado.save(ruta_completa_docx_modificado)

    print(f"Proceso completado. Resultado guardado en {ruta_completa_docx_modificado}")

if __name__ == "__main__":
    main()
