
from docx import Document
import openai
from pathlib import Path
import json
import os
import win32com.client

# Configuración de la clave de la API de OpenAI
API_KEY = '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'
openai.api_key = API_KEY

# Función para cargar y modificar el archivo JSON
def modificar_json(ruta_archivo):
    try:
        with open(ruta_archivo, 'r', encoding='utf-8') as archivo:
            data = json.load(archivo)

        if isinstance(data, list):
            # Filtrar las entradas con 'Numero_de_documento' no nulo o vacío
            data_filtrada = [registro for registro in data if isinstance(registro, dict) and registro.get('Numero_de_documento') not in [None, ""]]

            # Sobrescribir el archivo original con los datos filtrados y estructurados
            with open(ruta_archivo, 'w', encoding='utf-8') as archivo:
                json.dump(data_filtrada, archivo, ensure_ascii=False, indent=4)
                print("Json modificado")
        else:
            print("El contenido del archivo no es una lista de objetos JSON.")
    except Exception as e:
        print(f"Ocurrió un error: {e}")

def BuscarPalabra(file_path, target_words):
    # Abre el documento .docx usando python-docx
    doc = Document(file_path)

    for i, para in enumerate(doc.paragraphs):
        for target_word in target_words:
            if target_word in para.text:
                # Encuentra la página basada en el índice del párrafo
                page_number = i // 25 + 1  # Suponiendo 25 párrafos por página
                return target_word, page_number
    
    return None, None  # Return None if the word is not found in the documen

def segmentar_texto(texto, max_tokens=4000):
    # Dividir el texto en segmentos de acuerdo con el límite de tokens
    segmentos = []
    palabras = texto.split()
    segmento_actual = palabras[0]
    
    for palabra in palabras[1:]:
        if len(segmento_actual) + len(palabra) + 1 <= max_tokens:
            segmento_actual += ' ' + palabra
        else:
            segmentos.append(segmento_actual)
            segmento_actual = palabra
    
    segmentos.append(segmento_actual)
    return segmentos

def main():
    # Lista de palabras clave a buscar en el documento
    target_words = ["Datos de EL CLIENTESS", "DNI", "D.N.I", "Documento de identidad"]

    try:
        # Obtener una lista de archivos .docx en la carpeta documentos, excluyendo los archivos temporales
        archivos_docx = [archivo for archivo in os.listdir("documentos") if archivo.lower().endswith('.docx') and not archivo.startswith('~$')]
        print("Buscando documentos .docx")

        # Verificar si se encontraron archivos .docx en la carpeta
        if not archivos_docx:
            print("No se encontraron archivos .docx en la carpeta especificada.")
            return

        for archivo_docx in archivos_docx:
            # Obtener la ruta completa del archivo .docx
            documento_path = os.path.join("documentos", archivo_docx)

            # Buscar la palabra y obtener el número de página
            target_word, page_number = BuscarPalabra(documento_path, target_words)

            if target_word is not None and page_number is not None:
                print(f'La palabra "{target_word}" se encuentra en la página: {page_number}')

                # Leer el texto del documento
                doc = Document(documento_path)
                texto_documento = "\n".join([para.text for para in doc.paragraphs])

                # Segmentar el texto en trozos manejables
                segmentos = segmentar_texto(texto_documento)

                # Obtener el segmento de la página donde se encontró la palabra
                segmento_pagina = segmentos[page_number - 1]

                # Generar el JSON utilizando OpenAI GPT-3.5 Turbo (modelo de chat)
                prompt = '''Toma el siguiente texto y extrae de él una lista de objetos JSON de cada una de las personas que se encuentran allí. 
                Estos objetos deben tener las siguientes propiedades: 
                Generar un Id a cada uno, empieza por el 1.
                Tipo (natural o jurídica) 
                Nombre 
                Nacionalidad 
                Tipo_de_documento(DNI o RUC) 
                Numero_de_documento
                Estado_civil (casado, divorciado, viudo, soltero o sus equivalentes femeninos)
                Relaciones_conyugales
                Domicilio
                (Guarda los nombres de las propiedades tal cual tomando en cuenta las mayúsculas)
                Denominación que se le dará en el documento (después de 'a quien en adelante se le denominará') 
                Si existen representantes, deben ser una lista de enteros indicando la posición de cada representante en la lista de objetos.(Es importante identificad todo los representantes) 
                También debe haber un campo de relaciones para relaciones conyugales, que será una lista de objetos con propiedades: 
                Índice (corresponde al índice del objeto relacionado) 
                Tipo de relación (solo para relaciones maritales u otros tipos) 
                Las propiedades que no se encuentren deben ser NULL o vacías en el caso de propiedades de lista. 
                Por favor, responde solo con el archivo JSON solicitado, sin añadir contexto adicional, solo generar un array y no palabras fuera de las cadenas, evita decir "Aquí tienes el archivo JSON solicitado:" o similares frases'''  # (Tu prompt de OpenAI aquí)
                
                prompt_segmento = prompt + f"\n\nLa palabra '{target_word}' se encuentra en la página {page_number} del documento:\n\n{segmento_pagina}"

                print("Enviando solicitud a OpenAI para generar el JSON...")
                respuesta_openai = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[
                        {"role": "system", "content": "You are a helpful assistant."},
                        {"role": "user", "content": prompt_segmento}
                    ]
                )
                json_generado = respuesta_openai.choices[0].message["content"].strip()

                print("JSON generado por OpenAI:")
                print(json_generado)  # Imprime el JSON generado por OpenAI

                # Definir la ruta del archivo JSON de salida
                ruta_json = Path("json/archivo.json")

                # Guardar el JSON en el archivo
                with open(ruta_json, "w", encoding="utf-8") as archivo_json:
                    archivo_json.write(json_generado)

                print(f"JSON guardado en '{ruta_json}'.")

                # Llamar a la función para modificar el archivo JSON
                modificar_json(ruta_json)

            else:
                print(f'La palabras clave no se encontró en el documento.')

    except Exception as e:
        print(f"Ocurrió un error: {e}")

if __name__ == "__main__":
    main()
