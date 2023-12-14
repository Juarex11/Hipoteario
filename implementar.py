import datetime
from num2words import num2words
from docx import Document
from docx.shared import Pt
import json
import os

meses_en_espanol = {
    1: "ENERO",
    2: "FEBRERO",
    3: "MARZO",
    4: "ABRIL",
    5: "MAYO",
    6: "JUNIO",
    7: "JULIO",
    8: "AGOSTO",
    9: "SEPTIEMBRE",
    10: "OCTUBRE",
    11: "NOVIEMBRE",
    12: "DICIEMBRE"
}

def leer_json(ruta):
    with open(ruta, 'r', encoding='utf-8') as archivo_json:
        datos = json.load(archivo_json)
    return datos

def eliminar_espacios_finales(texto):
    return ''.join(texto.split())

def generar_detalles(registro, representante_id):
    detalles = []

    if registro["Id"]:
        # Verificar si el ID actual coincide con el representante_id
        if registro["Id"] == representante_id:
            detalles.append("REPRESENTADO POR")
        #detalles.append(f'ID: {registro["Id"]}')
        
    if registro["Nombre"]:
        detalles.append(f'Nombre: {registro["Nombre"]}')
    if registro["Tipo"]:
        detalles.append(f'Tipo: {registro["Tipo"]}')
    if registro["Nacionalidad"]:
        detalles.append(f'Nacionalidad: {registro["Nacionalidad"]}')
    if registro["Tipo_de_documento"]:
        detalles.append(f'Tipo de documento: {registro["Tipo_de_documento"]}')
    if registro["Numero_de_documento"]:
        detalles.append(f'Numero de documento: {registro["Numero_de_documento"]}')
    if registro["Estado_civil"]:
        detalles.append(f'Estado civil: {registro["Estado_civil"]}')
    if registro["Domicilio"]:
        detalles.append(f'Domicilio: {registro["Domicilio"]}')

    resultado = ', '.join(detalles)
    return resultado if resultado else None

def agregar_espaciado_entre_palabras(texto):
    return ' '.join(texto.split())

def cambiar_formato_texto(parrafo, fuente='Arial Narrow', tamaño=9):
    run = parrafo.runs[0]
    font = run.font
    font.name = fuente
    font.size = Pt(tamaño)
    run.text = run.text.upper()

def main():
    datos_json = leer_json('json/archivo.json')
    plantilla_doc = Document('plantilla/MODELO PLANTILLA.docx')

    fecha_actual = datetime.datetime.now()
    dia_en_letras = num2words(fecha_actual.day, lang='es').upper()
    mes_en_espanol = meses_en_espanol[fecha_actual.month]
    anio_en_letras = num2words(fecha_actual.year, lang='es').upper()
    texto_fecha = f"EL {dia_en_letras} DE {mes_en_espanol} DEL AÑO {anio_en_letras}"

    # Especifica el ID del representante que deseas imprimir
    representante_id = 3

    detalles = []

    for registro in datos_json:
        detalle = generar_detalles(registro, representante_id)
        if detalle is not None:
            detalles.append(detalle)

    resultado = '\n'.join(detalles)
    resultado = agregar_espaciado_entre_palabras(resultado)

    cadena_a_reemplazar = 'NOMBRE, PERUANO/A, SOLTERO/CASADO, OCUPACION …, CON DOCUMENTO NACIONAL DE IDENTIDAD …., CON DOMICILIO EN …..(NOTA VERIFICAR QUE SEA EL DEL D.N.I.; SINO INDICAR QUE EL DECLARA QUE SU DOMICILIO ACTUAL ES …. ); Y DE LA OTRA PARTE: NOMBRE, PERUANO/A, SOLTERO/CASADO, OCUPACION …, CON DOCUMENTO NACIONAL DE IDENTIDAD …., CON DOMICILIO EN …..'

    for p in plantilla_doc.paragraphs:
        if cadena_a_reemplazar in p.text:
            p.text = p.text.replace(cadena_a_reemplazar, resultado)
            cambiar_formato_texto(p)

    cadena_a_reemplazar_fecha = 'EL XXXXXXXXXXXXXX DE XXXXXXXXXXXXXX DEL AÑO DOS MIL XXXXXXXXXXXXXX'
    for p in plantilla_doc.paragraphs:
        if cadena_a_reemplazar_fecha in p.text:
            p.text = p.text.replace(cadena_a_reemplazar_fecha, texto_fecha)
            cambiar_formato_texto(p)

    if not os.path.exists('modificado'):
        os.makedirs('modificado')
    plantilla_doc.save('modificado/ejemplo_terminado.docx')
    print("Programa finalizado")

if __name__ == "__main__":
    main()
